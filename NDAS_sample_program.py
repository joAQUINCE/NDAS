# Import libraries
import numpy as np # Fundamental package for numerical computation (arrays, matrices, etc.).
import time # Provides time-related functions, used here for timing execution.
import sys # Provides access to system-specific parameters and functions.
import os # Provides a way of using operating system dependent functionality like reading/writing files.
from shutil import copyfile # Used for copying files, e.g., template documents.
from tqdm.notebook import tqdm # A fast, extensible progress bar for loops, optimized for Jupyter notebooks.
from scipy.spatial.transform import Rotation as R # Used for handling 3D rotations for orienting geometry.

# Import plotting libraries from Matplotlib
from mpl_toolkits.mplot3d import proj3d # Utilities for 3D projection.
from mpl_toolkits.mplot3d import Axes3D # Toolkits for creating 3D axes.

import matplotlib.colors as colors # For color mapping data values (e.g., stress levels).
import matplotlib.pyplot as plt # The main plotting interface for Matplotlib.
from matplotlib import cm # Colormaps module (e.g., cm.rainbow).

import matplotlib # Import the base matplotlib library.
matplotlib.rcParams['axes.linewidth'] = 0.2 # Set the default linewidth for plot axes globally.

# Import libraries for working with DOCX (Word) files
import docx # Main library for creating and modifying .docx files.
from docxcompose.composer import Composer # Library for combining multiple .docx files.

from docx.shared import Pt # Used to specify point sizes for fonts.
from docx.shared import RGBColor # Used to specify colors using RGB values.


# TO DO:
# Add limitation on the number of components (i.e. valves added to the coversheet)
# Developer note: A future enhancement is planned to limit the number of components listed on the coversheet.

# Copy template files from 'input_files' directory to 'runtime_files' directory.
# This prevents modification of the original templates during script execution.
copyfile("input_files/cover.docx", "runtime_files/cover.docx") # Copy the cover sheet template.
copyfile("input_files/nozzle_loads.docx", "runtime_files/nozzle_loads.docx") # Copy the nozzle loads template.

# Define the main class to handle the piping system analysis data.
class piping_system:

    # Constructor method for the class. Initializes instance variables.
    def __init__(self, output_filename, data_echo_deadweight_keywords, max_blank_line_cnt):

        # Store the path to the FEA output file.
        self.output_filename = output_filename

        # Store keywords used to identify relevant deadweight load cases in the output file.
        self.data_echo_deadweight_keywords = data_echo_deadweight_keywords

        # Initialize storage for parsed data (set to None initially, populated by methods).
        self.raw_data_list = None # Will hold all lines read from the output file.
        self.load_case_list = None # Will hold the names of identified load cases.
        self.loadcase_geometry_dic = {} # Dictionary to store geometry data (nodes, coordinates) per load case.
        self.loadcase_node_connectivity_dic = {} # Dictionary to store node connectivity (segments) per load case (unused later?).

        # Store the maximum number of consecutive blank lines allowed when searching for load cases.
        self.max_blank_line_cnt = max_blank_line_cnt

    # Method to parse the raw FEA output file line by line.
    def parse_raw_data(self):

        # Initialize an empty list to store each line read from the file.
        self.raw_data_list = []
        # Define the keyword string that identifies the start of the main stress analysis results section.
        stress_analysis_keyword = "ALL     STRESS ANALYSIS"

        # Define keywords related to the "Pipe Stress Summary" table within the output file.
        pipe_stress_summary_keyword = "    CONDITION      LEVEL     END    ELEMENT  STRESS(   PSI) (   PSI)  ALLOWABLE" # Header line.
        pipe_stress_summary_end_keyword = " APPROVED BY " # Keyword indicating the end of the summary section.
        
        # Initialize variables to store the line numbers where the summary starts and ends.
        self.pipe_stress_summary_start_line_number = None
        self.pipe_stress_summary_end_line_number = None

        # Define keywords related to the "INPUT CARD IMAGES" section, which often contains geometry definitions.
        input_card_start_keyword = "INPUT CARD IMAGES" # Marks the beginning of this section.
        input_card_end_keyword = "         .  +" # Marks the end of this section (specific format).
        
        # List of keywords (often 3-letter codes like 'TEA', 'END') found in specific columns that also signify the end of geometry input.
        end_geometry_keyword_list = ["TEA","ACE","END"]

        # Define delimiters used for parsing parameters embedded in comments (e.g., ***/ PARAM = VALUE).
        parameter_start_key = "***/" # Prefix indicating a parameter definition line.
        parameter_end_key = "="     # Separator between parameter name and value.
        comment_key = "***"       # General indicator for a comment line.

        # Define a list of parameter names expected to be found in comments, intended for the coversheet.
        self.coversheet_parameter_list = ["ANALYSIS NUMBER","ANALYSIS TITLE","STATION",\
                                        "UNIT NUMBER","DISCIPLINE","SAFETY CLASS",\
                                        "SYSTEM CODE","STRUCTURE","ANALYSIS REVISION",\
                                        "PACKAGE NUMBER","PACKAGE REVISION","AFFECTED DOCUMENT NUMBER",\
                                        "AFFECTED DOCUMENT REVISION",\
                                        "DOES ANALYSIS CONTAIN SAFEGUARDS INFORMATION(YER OR NO)?",\
                                        "DOES ANALYSIS CONTAIN UNVERIFIED ASSUMPTIONS(YER OR NO)?",\
                                        "UNVERIFIED ASSUMPTION TRACKING ORDER",\
                                        "SUPERSEDED DOCUMENT",\
                                        "PREPARER NAME","PREPARER SIGNATURE","DATE PREPARED",\
                                        "REVIEWER NAME", "REVIEWER SIGNATURE", "DATE REVIEWED",\
                                        "METHOD OF REVIEW","TYPE OF REVIEW","EXTERNAL APPROVER NAME",\
                                        "EXTERNAL APPROVER SIGNATURE","EXTERNAL APPROVAL DATE",\
                                        "COMPANY REVIEWER NAME","COMPANY REVIEWER SIGNATURE","COMPANY REVIEW DATE",\
                                        "INDEPENDENT THIRD PARTY REVIEW REQUIRED(YER OR NO)",\
                                        "COMPANY APPROVER NAME","COMPANY APPROVAL SIGNATURE","COMPANY APPROVAL DATE"]

        # Define parameters that might appear multiple times and should be collected into a list (e.g., multiple components).
        self.coversheet_set_list = ["INPUT DOC","FROM/TO","COMPONENT"]

        # Define parameters whose values might span multiple comment lines and need concatenation.
        self.long_text_parameter_list = ["DESCRIPTION OF CHANGE","DESCRIPTION OF REVISION"]

        # Define parameters whose content should be formatted differently in the final report (e.g., using Courier font for the stress summary).
        # Maps the parameter name to its intended section number in the report.
        self.courier_long_text_parameter_dic = {"PIPE STRESS SUMMARY":"5.1"}

        # Initialize a dictionary to store the extracted coversheet parameter values.
        self.coversheet_parameter_dic = {}

        # Define keywords typically found on input cards defining pipe segment properties.
        self.raw_segment_parameters_keyword_list = ["OD=","THI=","LBS/FT=","DPRESS=","PRESS=","PPRESS=",\
                                                   "E=","CODE=","CLASS=","SC=","SH=","SIF=","ADDWT=","ADD="] # (Outer Diameter, Thickness, etc.)

        # Define keywords relevant for nozzle analysis, often a subset of segment parameters.
        self.nozzle_analysis_keyword_list = ["OD=","THI="] # Primarily interested in OD and Thickness for nozzle calculations.

        # Define parameter names used specifically within the nozzle load check section/report template.
        self.nozzle_analysis_parameter_list = ["NOZZLE NODE POINT","ALLOWABLE AXIAL LOAD","ALLOWABLE RESULTANT SHEAR LOAD",\
                                                "ALLOWABLE TORSION MOMENT","ALLOWABLE BENDING MOMENT",\
                                                "OUTSIDE DIAMETER","OUTSIDE DIAMETER", # Duplicate entry?
                                                "NOZZLE THICKNESS","NOZZLE AREA","NOZZLE SECTION MODULUS",\
                                                "YIELD","K1","K2"] # Allowables, geometry, material/factors.

        # Define parameter names expected within the special "THOR" comment lines.
        self.THOR_parameter_list = ["THOR NODE","ME101 NODE NUMBER", "THOR NODE NUMBER", "NODE ELEVATION", "NODE OUTSIDE DIAMETER", "NODE INSIDE DIAMETER", "NONE PIPE SCHEDULE", "ORIGIN", "ORIGIN COORDINATES"]

        # Initialize a list to store parsed segment data (node pairs and their associated parameters).
        # Format: [[node_id1, {param1: val1, ...}], [node_id2, {param1: val1, ...}], ...] where each pair represents a segment.
        self.raw_segment_parameter_list = []
        # Initialize an empty dictionary (appears unused).
        parameters_old_dic = {}
        # Initialize an empty dictionary (appears unused, likely superseded by self.nozzle_load_parameter_dic).
        self.current_nozzle_analysis_parameters = {}
        # Flag to track whether the parser is currently inside the "INPUT CARD IMAGES" section.
        parsing_raw_geometry = False
        # Variable to hold endpoints (appears unused).
        segment_endpoints = None
        # Flag/State variable to handle single-node input cards correctly:
        # None = initial state, True = just processed a single-node start card, False = processed a two-node card or subsequent single node.
        first_single_node_in_this_line = None
        # List to track endpoints parsed from consecutive lines, used to detect changes. Initialized with a dummy value.
        segment_all_endpoints_list = [""]
        # Initialize an empty list (appears unused).
        segment_node_list = []

        # Temporary list for holding split THOR parameters read from a line.
        THOR_parameters_lst = []

        # Dictionary to store parsed THOR node information, keyed by the 'ME101 NODE NUMBER'.
        self.THOR_node_dic = {}
        # Flag to indicate if an error was encountered while parsing THOR parameters.
        self.THOR_input_error = False

        # Open the FEA output file for reading.
        # Uses 'cp1252' encoding, common for older Windows-based FEA software like ME101/ADLPIPE. Error handling might be needed if encoding varies.
        with open (self.output_filename,encoding='cp1252') as raw_data:
            # Loop through each line in the file, keeping track of the line number (line_cnt).
            for line_cnt,line in enumerate(raw_data):

                # Append the raw line (string) to the list holding the entire file content.
                self.raw_data_list.append(line)

                # --- THOR Parameter Parsing Logic ---
                # Check if the line contains the specific THOR parameter keyword 'THOR NODE' at a fixed position (columns 14-23).
                if line[14:23] == self.THOR_parameter_list[0]: # Check for "THOR NODE" keyword.
                    # Extract the parameter string assumed to be enclosed in curly braces {}.
                    THOR_parameters_str = line[line.find("{")+1:line.find("}")]

                    # Check if the essential 'ME101 NODE NUMBER' parameter exists within the extracted string and no previous THOR error occurred.
                    if self.THOR_parameter_list[1] in line and not self.THOR_input_error:
                        # Find the start position of the ME101 node number's value (after the '=').
                        node_number_start = THOR_parameters_str.find(self.THOR_parameter_list[1])+len(self.THOR_parameter_list[1])
                        # Find the end position of the node number's value (before the next comma).
                        node_number_end   = THOR_parameters_str[node_number_start:].find(",") + node_number_start
                        # Extract the ME101 node number string, remove '=' and strip whitespace.
                        node_number = THOR_parameters_str[node_number_start:node_number_end].replace("=","").strip()
                        # Extract the remaining parameter string after the node number.
                        THOR_parameters_lst = THOR_parameters_str[node_number_end +1:].split(",")
                        # Split each remaining parameter string (e.g., "KEY=VALUE") into a [key, value] pair.
                        THOR_parameters_lst = [i.split("=") for i in THOR_parameters_lst]

                        # Try to create a dictionary from the key-value pairs for this node.
                        try:
                            # Store the dictionary of parameters in self.THOR_node_dic, keyed by the ME101 node number.
                            self.THOR_node_dic[node_number] = {i[0].strip():i[1].strip() for i in THOR_parameters_lst}
                        # If splitting or dictionary creation fails (e.g., bad format)...
                        except:
                            # Print a basic error indicator. Consider more descriptive logging.
                            print("BA") # Indicates an error during THOR parameter dictionary creation.
                            # Set the flag indicating a THOR parsing error occurred.
                            self.THOR_input_error = True
                    # If the 'ME101 NODE NUMBER' was missing in a line identified as a THOR line...
                    else:
                        # Print a different basic error indicator.
                        print("B") # Indicates missing 'ME101 NODE NUMBER' in a THOR line.
                        # Set the flag indicating a THOR parsing error occurred.
                        self.THOR_input_error = True

                # --- Section Boundary Identification ---
                # Check if the line contains the keyword marking the start of the main stress analysis section.
                if stress_analysis_keyword in line:
                    # Store the line number where this section begins.
                    self.stress_analysis_line_number = line_cnt

                # Check if the line contains the specific header for the pipe stress summary table.
                if pipe_stress_summary_keyword in line:
                    # Store the line number where the summary table header is found.
                    self.pipe_stress_summary_start_line_number = line_cnt

                # If the start of the summary has been found, check if the current line contains the keyword marking its end.
                if self.pipe_stress_summary_start_line_number != None and pipe_stress_summary_end_keyword in line:
                    # Store the line number where the summary section ends.
                    self.pipe_stress_summary_end_line_number = line_cnt

                # --- Coversheet Parameter Parsing (from comments) ---
                # Check if the line starts with the designated parameter prefix ('***/').
                if parameter_start_key in line:
                    # Find the start index of the parameter name (after the prefix).
                    parameter_name_start_character  = line.find(parameter_start_key)+len(parameter_start_key)
                    # Find the end index of the parameter name (at the '=' sign).
                    parameter_name_end_character    = line.find(parameter_end_key)
                    # Find the start index of the parameter value (after the '=' sign).
                    parameter_value_start_character = line.find(parameter_end_key) + 1
                    # Find the end index of the parameter value (end of the line).
                    parameter_value_end_character   = len(line)
                    # Extract the parameter name (key) and strip whitespace.
                    parameter_key = line[parameter_name_start_character:parameter_name_end_character].strip()

                    # Check if the extracted key is in the list of desired coversheet parameters.
                    if parameter_key in self.coversheet_parameter_list:
                        # Store the parameter key and its extracted value (stripped) in the coversheet dictionary.
                        self.coversheet_parameter_dic[parameter_key] = line[parameter_value_start_character:parameter_value_end_character].strip()

                    # Check if the extracted key is in the list of parameters expected to have multiple values.
                    if parameter_key in self.coversheet_set_list:
                        # If this key hasn't been seen before...
                        if parameter_key not in self.coversheet_parameter_dic:
                            # Initialize an empty list for this key in the dictionary.
                            self.coversheet_parameter_dic[parameter_key] = []
                        # Append the current value (stripped) to the list associated with this key.
                        self.coversheet_parameter_dic[parameter_key].append(line[parameter_value_start_character:parameter_value_end_character].strip())

                    # Check if the extracted key is in the list of parameters whose values might span multiple lines.
                    if parameter_key in self.long_text_parameter_list:
                        # If this key hasn't been seen before...
                        if parameter_key not in self.coversheet_parameter_dic:
                            # Initialize an empty string for this key in the dictionary.
                            self.coversheet_parameter_dic[parameter_key] = str()
                        # Append the current line's value (stripping newline characters) to the existing string for this key.
                        self.coversheet_parameter_dic[parameter_key] += line[parameter_value_start_character:parameter_value_end_character].replace("\n","")

                # --- Geometry Parsing (within INPUT CARD IMAGES section) ---
                # This block executes only if the 'parsing_raw_geometry' flag is True.
                if parsing_raw_geometry:

                    # Check for conditions that indicate the end of the geometry input section.
                    # Condition 1: Specific end keyword '         .  +' is found.
                    # Condition 2: A 3-letter keyword from end_geometry_keyword_list ('TEA', 'ACE', 'END') is found in columns 10-13.
                    if input_card_end_keyword in line or line[10:13] in end_geometry_keyword_list:
                        # If end condition is met, set the flag to False to stop geometry parsing.
                        parsing_raw_geometry = False
                    # If not an end-of-geometry line, proceed to parse potential geometry/parameter data.
                    else: # Added 'else' for clarity, though the 'if' above would prevent falling through if True.
                        # Check if the line likely defines a segment (based on non-empty 'FROM'/'TO' columns and not being a comment/end keyword).
                        # Assumes 'FROM' node is cols 13-16, 'TO' node is cols 16-19. This is typical for ADLPIPE/ME101.
                        if line[13:19].replace(" ","") != "" and comment_key not in line: # Removed redundant check for end_geometry_keyword_list

                            # Extract potential 'FROM' and 'TO' node IDs from their fixed columns.
                            segment_endpoints_list = [line[13:16].strip(),line[16:19].strip()]

                            # Append the currently parsed endpoints to the list tracking endpoint history.
                            segment_all_endpoints_list.append(segment_endpoints_list)

                            # Check if the endpoints on this line are different from the endpoints on the immediately preceding line.
                            # This detects the start of a new segment definition line (or the very first line).
                            if segment_all_endpoints_list[-1] != segment_all_endpoints_list[-2]:

                                # --- Handle Different Input Card Formats (Single vs. Double Node) ---
                                # Check if only one node ID is present on the line (i.e., 'FROM' is blank, 'TO' is present).
                                # Also check if this is the *very first* node definition encountered (first_single_node_in_this_line is None).
                                if len([i for i in segment_endpoints_list if i != ""]) == 1 and first_single_node_in_this_line == None:

                                    # If it's the first single node, update the initial dummy entry in segment_all_endpoints_list.
                                    segment_all_endpoints_list[0] = line[16:19].strip() # Store the actual first node ID.

                                    # Add the first node ID and an associated empty dictionary (for parameters) to the main segment list.
                                    self.raw_segment_parameter_list.append([line[16:19].strip(),{}])

                                    # Set the state flag to indicate that a single-node start card was just processed.
                                    first_single_node_in_this_line = True

                                # Check if only one node ID is present, but it's *not* the first node overall.
                                # This implies a continuation card, where the 'FROM' node is implicitly the 'TO' node of the previous segment.
                                elif len([i for i in segment_endpoints_list if i != ""]) == 1:

                                    # Get the 'FROM' node ID from the *last* node entry added to the segment list.
                                    from_node = self.raw_segment_parameter_list[-1][0]

                                    # Extend the segment list with two entries:
                                    # 1. The inferred 'FROM' node (using the ID from the previous entry) with an empty parameter dict.
                                    # 2. The new 'TO' node (from the current line) with an empty parameter dict.
                                    
                                    # NOTE: This structure seems to explicitly create pairs, duplicating the 'from_node' ID reference. Review if this is the desired structure or if only the 'to_node' should be added here.
                                    self.raw_segment_parameter_list.extend([[from_node,{}],[line[16:19].strip(),{}]])

                                    # Set the state flag to False, indicating a segment connection (implicit or explicit) was made.
                                    first_single_node_in_this_line = False

                                # Handle the case where *both* 'FROM' and 'TO' nodes are specified on the line.
                                else:
                                    # Extend the segment list with two entries:
                                    # 1. The 'FROM' node from the current line with an empty parameter dict.
                                    # 2. The 'TO' node from the current line with an empty parameter dict.
                                    self.raw_segment_parameter_list.extend([[line[13:16].strip(),{}],[line[16:19].strip(),{}]])

                                    # Set the state flag to False.
                                    first_single_node_in_this_line = False

                        # --- Parse Segment Parameters (OD, THI) associated with the nodes just defined ---
                        # Iterate through the list of keywords defining segment properties (currently just OD=, THI=).
                        for keyword in self.nozzle_analysis_keyword_list: # Using nozzle list as it contains OD, THI
                            # Check if the keyword exists on the current line and it's not a comment line.
                            if keyword in line and comment_key not in line:

                                # Find start index of the parameter value (after the keyword).
                                start = line.find(keyword)+len(keyword)
                                # Find the end index of the value (at the next comma, or fallback).
                                end = line[start:].find(",")
                                # If no comma is found, assume the value extends to a fixed column (92). This might need adjustment based on file format.
                                if end == -1:
                                    end = 92 # Fallback end position if comma is not found.
                                else:
                                    end += start # Adjust end position relative to the start of the line.

                                # Extract the parameter value, remove spaces, and strip whitespace.
                                parameter_value = line[start:end].replace(" ","").strip() # Extract between start and calculated end.

                                # --- Assign Parameter Value to Nodes ---
                                # If the current line defined the *first single node*:
                                if first_single_node_in_this_line == True: # Explicit check for True
                                    # Add the parameter (key without '=') and its value to the dictionary of the *last* node added (the single start node).
                                    self.raw_segment_parameter_list[-1][1][keyword[:-1]] = parameter_value

                                # If the current line defined a segment (two nodes explicitly, or one node continuing from previous):
                                elif first_single_node_in_this_line == False: # Explicit check for False
                                    # Add the parameter to the dictionary of the 'FROM' node of the segment (second to last entry).
                                    self.raw_segment_parameter_list[-2][1][keyword[:-1]] = parameter_value
                                    # Add the *same* parameter to the dictionary of the 'TO' node of the segment (last entry).
                                    # Assumes the parameter applies to the entire segment defined by these nodes.
                                    self.raw_segment_parameter_list[-1][1][keyword[:-1]] = parameter_value

                                # --- Handle Parameter Inheritance (Potentially problematic logic) ---
                                # This 'elif' condition seems complex and might be redundant or handle edge cases where parameters
                                # are defined on lines *after* the node definition without repeating the nodes.
                                # It checks: not a comment, not the first single node state, AND the keyword is NOT already in the last node's dictionary.
                                elif comment_key not in line and first_single_node_in_this_line == False \
                                and keyword[:-1] not in self.raw_segment_parameter_list[-1][1]:
                                    # If these conditions are met, it tries to inherit the parameter value from the 'FROM' node of the *previous* segment.
                                    # This assumes a specific structure where parameters might carry over if not redefined.
                                    # Accessing [-3] could lead to index errors if called too early.
                                    try: # Add try-except for safety
                                        parameter_value = self.raw_segment_parameter_list[-3][1][keyword[:-1]]
                                        # Apply the inherited value to the 'FROM' node of the *current* segment.
                                        self.raw_segment_parameter_list[-2][1][keyword[:-1]] = parameter_value
                                        # Apply the inherited value to the 'TO' node of the *current* segment.
                                        self.raw_segment_parameter_list[-1][1][keyword[:-1]] = parameter_value
                                    except IndexError:
                                        print(f"Warning: Could not inherit parameter '{keyword[:-1]}' due to insufficient history.")
                                    except KeyError:
                                         print(f"Warning: Could not inherit parameter '{keyword[:-1]}' as it was missing in the previous segment.")


                # --- Start Geometry Parsing Trigger ---
                # Check if the line contains the keyword indicating the start of the input card images section.
                if input_card_start_keyword in line:
                    # If found, set the flag to True, enabling the geometry parsing logic block for subsequent lines.
                    parsing_raw_geometry = True

    # Method to parse and identify the load case names from the raw data.
    def parse_load_cases(self):
            # Ensure raw data has been parsed first. If not, call parse_raw_data().
            if self.raw_data_list == None:
                self.parse_raw_data()

            # Define the keyword used to identify lines defining load cases (e.g., "LDCASE=WTDW(").
            load_case_keyword = "LDCASE="
            # Define the character that typically follows the load case name.
            load_case_keyword_end = "("

            # Initialize an empty list to store the names of the identified load cases.
            self.load_case_list = []
            # Flag to indicate whether the parser is currently within the block of load case definitions.
            loading_loadcases = False
            # Counter for consecutive lines without a load case definition after the first one was found.
            no_loadcase_line_cnt = 0

            # Iterate through each line stored in the raw_data_list.
            for line in self.raw_data_list:

                # Check if the consecutive non-loadcase line count exceeds the maximum allowed limit.
                if no_loadcase_line_cnt > self.max_blank_line_cnt:
                    # If exceeded, assume all load cases have been found and break the loop.
                    break

                # Check if the current line contains the load case definition keyword.
                if load_case_keyword in line:
                    # Extract the load case name found between the keyword and the end character.
                    load_case_name = line[line.find(load_case_keyword)+len(load_case_keyword): \
                                         line.find(load_case_keyword_end)]
                    # Add the extracted load case name to the list.
                    self.load_case_list.append(load_case_name)
                    # Set the flag indicating that load case definitions are being processed.
                    loading_loadcases = True
                    # Reset the non-loadcase line counter as a load case was found.
                    no_loadcase_line_cnt = 0 # Reset counter when a load case is found

                # If currently processing load cases, but the current line does *not* contain the keyword...
                elif loading_loadcases: # Use elif to avoid incrementing on the same line a load case was found
                    # Increment the counter for consecutive lines without load case definitions.
                    no_loadcase_line_cnt += 1

            # Print a confirmation message and the list of found load cases.
            print("Load Case Parsing Complete\nLoad Case List:", self.load_case_list)

    # Method to check if the parsed node coordinates are consistent across different load cases.
    def get_cordinate_distance(self): # Typo: should be get_coordinate_distance
        # Ensure node coordinates have been parsed first.
        self.parse_node_coordinates() # Assumes this populates self.loadcase_geometry_dic

        # Initialize a list to store the calculated distances (norms of differences) between coordinate sets.
        distance_check_list = []
        # Iterate through all pairs of load cases found.
        for key1 in self.loadcase_geometry_dic:
            for key2 in self.loadcase_geometry_dic:
                # Calculate the Euclidean distance (norm) between the coordinate arrays (excluding the node ID column) for the two load cases.
                # Assumes coordinates are stored from the second column onwards and need conversion to float.
                distance = np.linalg.norm(self.loadcase_geometry_dic[key1][:,1:].astype(float) - self.loadcase_geometry_dic[key2][:,1:].astype(float))
                # Append the calculated distance to the list.
                distance_check_list.append(distance)

        # Check if the sum of all calculated distances is effectively zero (within floating-point tolerance).
        if sum(distance_check_list) < 1e-9: # Use a small tolerance instead of exact zero
            # If the sum is zero, it implies all load cases have identical node coordinates.
            print("Distance Check Completed Successfully. Loadcases parse identical geometries.")
        else:
            # If the sum is non-zero, it indicates discrepancies in node coordinates between load cases.
            print("Distance Check Failed. Different Loadcases parse different geometries.")

    # Method to parse node connectivity information (which nodes form pipe segments).
    def get_node_connectivities(self):

        # Ensure the raw data has been parsed (needed to find the stress analysis section).
        if self.raw_data_list == None:
            self.parse_raw_data() # Should ideally call parse_load_cases if not already called

        # Flag to control when to start parsing connectivity data.
        parse_data = False
        # Initialize list to store the pairs of connected node IDs.
        node_connectivity_list = []
        # Counter for skipped lines (unused in current logic).
        skipped_line_cnt = 0
        # Initialize the start line number for connectivity parsing to a very large value.
        node_connectivity_start_line = 10**100 # Represents infinity initially

        # Keyword indicating the header line just before the connectivity data starts.
        node_connectivity_keyword = " FROM "
        # Get the portion of the raw data starting from the stress analysis section.
        # This assumes connectivity is listed within or after the stress analysis.
        node_stress_list = self.raw_data_list[self.stress_analysis_line_number:]

        # Iterate through the lines in the relevant section of the output file.
        for line_cnt,line in enumerate(node_stress_list):
            # Check if the line contains the keyword identifying the connectivity header.
            if node_connectivity_keyword in line:
                # Store the line number (relative to node_stress_list) where the header is found.
                node_connectivity_start_line = line_cnt

            # Check if the current line number is exactly 4 lines after the header.
            # This assumes a fixed offset between the header and the start of data. This is brittle.
            if line_cnt == node_connectivity_start_line + 4:
                # Enable the data parsing flag.
                parse_data = True

            # --- Parse Connectivity Data ---
            # Check if parsing is enabled AND if specific columns (25-115) on the current and next lines contain numeric data (or '*').
            # This is a very specific format check, likely tied to ME101/ADLPIPE stress output format.
            if parse_data and \
               line[25:115].replace(" ","").replace("*","").replace(".","").isnumeric() and \
               (line_cnt + 1 < len(node_stress_list)) and \
               node_stress_list[line_cnt+1][25:115].replace(" ","").replace("*","").replace(".","").isnumeric(): # Added check for next line existence and decimal points
                # If conditions are met, extract the node ID from the current line (cols 0-8) and the next line (cols 0-8).
                # These form a connected pair (segment). Strip whitespace from IDs.
                node_connectivity_list.append([line[0:8].strip(),node_stress_list[line_cnt+1][0:8].strip()])

            # Check if the line contains text indicating the end of the relevant section.
            if parse_data and "PIPING STRESS SUMMARY CHECK AND COVER SHEET" in line:
                # Disable the data parsing flag.
                parse_data = False

        # Convert the list of node pairs into a NumPy array for easier manipulation.
        self.node_connectivity_array = np.asarray(node_connectivity_list)
        # print(f"Node Connectivity Array Shape: {self.node_connectivity_array.shape}") # Debug print

    # Method to parse node coordinates from the "NODE DATA" section for each load case.
    def parse_node_coordinates(self):
        # Initialize a list to store all unique coordinate node IDs encountered.
        self.all_coordinate_id_list = []
        # Ensure load cases have been identified first.
        if self.load_case_list == None:
            self.parse_load_cases()

        # Keyword identifying the start of the node data section in the output.
        data_echo_keyword = "NODE DATA"
        # Flag to control when to parse node data.
        parse_data = False
        # Initialize a dictionary to store geometry data, keyed by load case name. Each value will be a list of tuples.
        geometry_dic  = {key:[] for key in self.load_case_list}
        # Variable to keep track of the current load case being parsed.
        current_load_case = None # Initialize

        # Iterate through each line in the raw data list.
        for line in self.raw_data_list:

            # Check if the line marks the start of a "NODE DATA" section for a known load case.
            # Assumes load case name appears in columns 15-24.
            if data_echo_keyword in line and line[15:24].strip() in self.load_case_list:
                # Store the current load case name.
                current_load_case = line[15:24].strip()
                # Enable the node data parsing flag.
                parse_data = True

            # Check if the line marks the end of the node data section (start of "ELEMENT DATA").
            if parse_data and "ELEMENT DATA" in line:
                # Disable the node data parsing flag.
                parse_data = False

            # If parsing node data, check if the coordinate columns (34-61) contain valid numeric data (allowing '-', '.').
            # This is a fixed-column format check.
            if parse_data and line[34:61].replace("-","").replace(".","").replace(" ","").isnumeric() and current_load_case is not None:
                # Extract the node ID (cols 0-8).
                node_id_raw = line[0:8].strip()
                # Extract coordinates X (cols 33-39), Y (44-50), Z (55-61).
                x_coord = line[33:39]
                y_coord = line[44:50]
                z_coord = line[55:61]

                # --- Node ID Matching Logic (Handle potential ' M' suffix) ---
                # Check if the raw node ID is NOT directly in the connectivity list, but the ID followed by " M" IS.
                # This handles potential discrepancies where connectivity uses a modified ID (e.g., for midpoints).
                # This check relies on self.node_connectivity_array being populated correctly.
                if node_id_raw not in self.node_connectivity_array[:,0] and \
                   node_id_raw not in self.node_connectivity_array[:,1] and \
                   (node_id_raw + " M") in self.node_connectivity_array: # Check both columns? Or just flatten? Let's assume it can be in either.
                       # If the ' M' version exists in connectivity, use that ID for storage.
                       node_id_to_store = node_id_raw + " M"
                       # Append the node data (ID with ' M', X, Y, Z) to the list for the current load case.
                       geometry_dic[current_load_case].append((node_id_to_store, x_coord, y_coord, z_coord))
                # Otherwise (node ID found directly, or neither version found)...
                else:
                    # Use the node ID as extracted (cols 0-6, slightly different from above?).
                    node_id_to_store = line[0:6].strip() # Using 0:6 here - consistent?
                    # Append the node data (ID, X, Y, Z) to the list for the current load case.
                    geometry_dic[current_load_case].append((node_id_to_store, x_coord, y_coord, z_coord))
                    # Add the cleaned node ID (without ' M') to the list of all coordinate IDs.
                    self.all_coordinate_id_list.append(node_id_to_store.replace(" ","")) # Remove spaces

        # After processing all lines, convert the lists of tuples in geometry_dic into NumPy arrays.
        # This loop might overwrite if multiple load cases exist but only the last one's data is kept before conversion.
        # It should ideally be inside the main loop or handle accumulation differently.
        # Let's assume current_load_case holds the last valid one processed if multiple exist.
        # Correction: Need to convert *all* valid load cases found.
        # geometry_dic[current_load_case] = np.asarray(geometry_dic[current_load_case]) # Original line - incorrect placement
        # Corrected loop to convert all populated lists to arrays:
        for key in self.load_case_list:
             if key in geometry_dic and len(geometry_dic[key]) > 0: # Check if key exists and list is not empty
                 geometry_dic[key] = np.asarray(geometry_dic[key])


        # Filter the geometry dictionary to only include load cases that actually had coordinate data parsed.
        self.loadcase_geometry_dic = {} # Re-initialize to store only valid ones
        for key in geometry_dic:
            # Check if the entry exists and the array is not empty.
            if isinstance(geometry_dic[key], np.ndarray) and (geometry_dic[key].shape[0]) != 0:
                # Add the valid load case geometry (as a NumPy array) to the final dictionary.
                self.loadcase_geometry_dic[key] = geometry_dic[key]

        # Return the dictionary containing geometry data per load case.
        return self.loadcase_geometry_dic

    # Method to rotate/transform node coordinates based on plant-specific conventions.
    def axes_rotation(self,loadcase,plant_name):
        # Check if the specified plant name is "hc".
        if plant_name == "hc":
            # Extract the node IDs (first column) from the geometry data for the given load case. Add a new axis for hstack.
            node               =  self.loadcase_geometry_dic[loadcase][:,0][:,np.newaxis]
            # Extract original X coordinates (col 1), convert to float, scale by 1, add new axis. This becomes the new X.
            array_modified_x =  1 * self.loadcase_geometry_dic[loadcase][:,1].astype(float)[:,np.newaxis]
            # Extract original Z coordinates (col 3), convert to float, scale by -1, add new axis. This becomes the new Y.
            array_modified_y = -1 * self.loadcase_geometry_dic[loadcase][:,3].astype(float)[:,np.newaxis]
            # Extract original Y coordinates (col 2), convert to float, scale by 1, add new axis. This becomes the new Z.
            array_modified_z =  1 * self.loadcase_geometry_dic[loadcase][:,2].astype(float)[:,np.newaxis]
            # This specific transformation (X=X, Y=-Z, Z=Y) likely maps FEA coordinates (e.g., Y=up) to plant coordinates (e.g., Z=up, Y=North/South).

        # Horizontally stack the node IDs and the transformed X, Y, Z coordinates into a single array.
        return np.hstack([node,array_modified_x,array_modified_y,array_modified_z])

    # Method to parse stress values from the stress analysis section.
    def parse_stresses(self):
        # Initialize a list to store sustained stress data [node_id, stress_value, allowable_stress?].
        sustained_stress_list =[]
        # Define the starting line number for parsing (previously stored).
        start = self.stress_analysis_line_number
        # Iterate through the lines in the raw data starting from the stress analysis section.
        for line_cnt, line in enumerate(self.raw_data_list[start:]):
            # Check if specific columns (18-130) contain numeric data (or '*'). This is a format-specific check for stress result lines.
            # Also check if the line is long enough to avoid index errors.
            if len(line) > 130 and line[18:130].replace("*","").replace(" ","").replace(".","").isnumeric(): # Added decimal point check
                # Extract the node ID from columns 0-7.
                node_id = line[0:7].strip()

                # --- Node ID Matching Logic (similar to coordinate parsing) ---
                # Check if the extracted node ID is NOT in the list of coordinate IDs,
                # BUT the ID without its last character IS in the list.
                # This might handle cases like '10A' vs '10' if coordinates only store '10'. This seems fragile.
                if node_id not in self.all_coordinate_id_list and \
                   node_id[:-1] in self.all_coordinate_id_list:
                   # If the shorter ID exists, print a warning/notification.
                   print("NOT FOUND, using shorter ID:",node_id) # Modified print message
                   # Use the shorter node ID for storing the stress.
                   node_id = node_id[:-1]

                # Append the potentially modified node ID and the stress value (cols 28-35) and allowable stress? (cols 35-44) to the list.
                # Assumes these columns contain the relevant stress values.
                sustained_stress_list.append([node_id,line[28:35].strip(),line[35:44].strip()])

        # Convert the list of stress data into a NumPy array.
        self.sustained_stresses = np.asarray(sustained_stress_list)
        # print(f"Sustained Stresses Array Shape: {self.sustained_stresses.shape}") # Debug print

    # Method to generate a 3D stress contour plot for a given load case.
    def get_stress_plot(self,loadcase):

        # Get the node coordinates for the specified load case, applying the plant-specific rotation.
        coordinates = self.axes_rotation(loadcase,"hc") # Hardcoded plant name "hc"

        # Create a Matplotlib figure and axes object. Adjust figure size.
        fig1, _ = plt.subplots(1, 1,figsize=(22,17))

        # Define a grid specification for arranging subplots within the figure. 100x100 grid.
        gs = fig1.add_gridspec(100, 100)

        # Initialize a list to store surface objects created by plot_surface (for potential future use, maybe transparency).
        surf_list = []

        # Add the main 3D subplot to the figure, spanning a large portion of the grid.
        ax1 = fig1.add_subplot(gs[1:-20,1:-20], projection='3d') # Leave space at bottom and right

        # Turn off the display of the 3D axes lines and labels for a cleaner look.
        ax1._axis3don = False

        # Generate a random 3D unit vector. Used later to define the plane for cylinder generation.
        random_vec  = np.random.rand(3)
        random_vec /= np.linalg.norm(random_vec) # Normalize to unit length.

        # Print status message.
        print(f"Processing Stress Contour Plot for {loadcase} Loadcase...")
        # Iterate through each segment defined in the node connectivity array, using tqdm for a progress bar.
        for segment_cnt in tqdm(range(len(self.node_connectivity_array))):

            # Get the current segment's start and end node IDs.
            segment = self.node_connectivity_array[segment_cnt]

            # --- Get Coordinates and Stresses for the Segment ---
            # Find the coordinates of the origin node in the rotated coordinate array. Squeeze removes extra dimensions. Convert to float. Scale by 1.3 (?).
            origin      = np.squeeze(coordinates[np.where(coordinates[:,0] == segment[0])])[1:].astype(float)*1.3 # Why 1.3 scaling? Visual separation?
            # Find the coordinates of the destination node. Squeeze, convert to float. Scale by 1.3.
            destination = np.squeeze(coordinates[np.where(coordinates[:,0] == segment[1])])[1:].astype(float)*1.3

            # Find the index of the origin node in the sustained stresses array.
            origin_node_array_index      = np.where(self.sustained_stresses[:,0] == segment[0].strip())[0][0] # Assumes unique node IDs
            # Find the index of the destination node in the sustained stresses array.
            destination_node_array_index = np.where(self.sustained_stresses[:,0] == segment[1].strip())[0][0] # Assumes unique node IDs

            # Set the font size for node labels.
            node_font_size = 3

            # --- Plot Node Labels ---
            # If this is the very first segment, plot the label for the origin node.
            if segment_cnt == 0:
                # Define properties for the node label bounding box (circle).
                bbox_props = dict(boxstyle="circle,pad=0.2", fc="w", ec="k", lw=.5)
                # Add the text label for the origin node at its coordinates.
                t = ax1.text(origin[0], origin[1], origin[2], segment[0], ha="center", va="center", rotation=0, \
                             size= node_font_size, \
                             bbox=bbox_props,zorder=100) # High zorder to be on top.

            # Plot a simple line representing the centerline of the pipe segment (optional, as surface will cover it).
            # ax1.plot([origin[0],destination[0]], [origin[1],destination[1]], [origin[2],destination[2]], color = 'b',lw=.3) # Commented out as surface is plotted

            # Prepare the label text for the destination node. Remove spaces if it's a modified ID (like "50 M").
            if len(str(segment[1])) > 3: # Simple check for likely modified IDs.
                label = str(segment[1]).replace(" ","")
            else:
                label = str(segment[1])

            # Define properties for the destination node label bounding box.
            bbox_props = dict(boxstyle="circle,pad=0.2", fc="w", alpha=0.8, ec="k", lw=.5) # Semi-transparent background
            # Add the text label for the destination node, slightly offset vertically based on segment count for less overlap.
            t = ax1.text(destination[0], destination[1]+.03*(-1)**segment_cnt, destination[2], label,\
                         ha="center", va="center", rotation=0,\
                         size= node_font_size, bbox=bbox_props,zorder=100)

            # --- Generate Cylinder Surface for the Segment ---
            # Define the number of points around the circumference and along the length of the cylinder.
            circ_point_count = 20 # Points around circumference.
            len_point_count = 10  # Points along the length.

            # Calculate the vector representing the segment from origin to destination.
            segment_vector   = destination - origin
            # Calculate the unit vector along the segment axis.
            segment_unit_vec = segment_vector/np.linalg.norm(segment_vector)

            # --- Calculate a Normal Vector for Cylinder Radius ---
            # Create a vector perpendicular to the segment axis using the random vector defined earlier.
            # Uses Gram-Schmidt orthogonalization concept: subtract projection of random_vec onto segment_unit_vec.
            random_normal_vec = random_vec - random_vec.dot(segment_unit_vec) * segment_unit_vec
            # Normalize the resulting normal vector and scale it to define the cylinder radius.
            # Division by 24 results in radius = 1/24 (units?). This controls the visual thickness of the pipe.
            random_normal_vec = random_normal_vec/np.linalg.norm(random_normal_vec)/24

            # --- Generate Points Around Circumference ---
            # Define angles for points around the circumference (0 to 2*pi).
            rotation_radians = np.linspace(0,2*np.pi,circ_point_count)
            # The axis of rotation is the segment's unit vector.
            rotation_axis = segment_unit_vec
            # Initialize array to store the vectors pointing from centerline to circumference points.
            segment_normal_vector_array = np.zeros([circ_point_count,3])

            # Rotate the initial normal vector around the segment axis to generate points on the circumference.
            for angle_cnt, angle in enumerate(rotation_radians):
                # Create rotation vector (axis-angle representation).
                rotation_vector = rotation_radians[angle_cnt] * rotation_axis
                # Create a SciPy Rotation object from the rotation vector.
                rotation = R.from_rotvec(rotation_vector)
                # Apply the rotation to the initial normal vector to get a vector for this angle.
                segment_normal_vector_array[angle_cnt] = rotation.apply(random_normal_vec)

                # Calculate a point on the circumference at the origin (scaled by 0.2?). Seems like debug code.
                origin_plus_unit_vec = origin + segment_normal_vector_array[angle_cnt]*.2

                # --- Debug Plotting (for node "60") ---
                # If the destination node ID is "60", plot helper lines (likely for debugging orientation).
                if self.sustained_stresses[destination_node_array_index,0] == "60":
                    # Points for plotting the scaled normal vector at the origin.
                    x_test = [origin[0],origin_plus_unit_vec[0]]
                    y_test = [origin[1],origin_plus_unit_vec[1]]
                    z_test = [origin[2],origin_plus_unit_vec[2]]
                    # Points for plotting the segment vector itself from the origin.
                    x_test1 = [origin[0],origin[0]+segment_vector[0]] # Should be destination[0]
                    y_test1 = [origin[1],origin[1]+segment_vector[1]] # Should be destination[1]
                    z_test1 = [origin[2],origin[2]+segment_vector[2]] # Should be destination[2]
                    # Plot the segment vector line in magenta.
                    ax1.plot(x_test1,y_test1,z_test1,c='m') # Plots the segment centerline again if node is 60

            # --- Generate Points Along Cylinder Length ---
            # Create points along the centerline of the segment.
            segments_discrete = np.linspace(origin,destination,len_point_count)
            # Initialize array to store all points on the cylinder surface.
            all_segment_points = np.empty([len_point_count,circ_point_count,3])

            # Generate the grid of points on the cylinder surface.
            for discrete_segment_cnt in range(len_point_count):
                # For each point along the centerline, add the array of normal vectors (scaled appropriately earlier)
                # to generate the ring of points at that position along the length.
                all_segment_points[discrete_segment_cnt] = \
                    segments_discrete[discrete_segment_cnt] + segment_normal_vector_array

            # --- Prepare Data for Surface Plotting ---
            # Extract X, Y, Z coordinates for the surface plot. Transpose needed for plot_surface format.
            x = np.transpose(all_segment_points[:,:,0])
            y = np.transpose(all_segment_points[:,:,1])
            z = np.transpose(all_segment_points[:,:,2])

            # Get the stress values at the origin and destination nodes. Convert to float.
            origin_stress      = self.sustained_stresses[origin_node_array_index,1].astype(float)
            destination_stress = self.sustained_stresses[destination_node_array_index,1].astype(float)

            # Linearly interpolate stress values along the length of the segment.
            C_dat   = np.linspace(origin_stress,destination_stress,len_point_count)

            # Determine the overall min and max stress values across all nodes for consistent color mapping.
            C_range = np.array([np.amin(self.sustained_stresses[:,1].astype(float)),\
                                np.amax(self.sustained_stresses[:,1].astype(float))])

            # Create a Normalize instance to map stress values to the range [0, 1] for the colormap.
            norm = colors.Normalize(vmin=C_range.min(),vmax=C_range.max())

            # Create a 2D array to hold the color data for the surface plot.
            # Dimensions match the surface grid (circumference points x length points).
            ci = np.zeros([len_point_count,circ_point_count]) # Initialize with zeros.

            # Fill the color array. Each column (representing a step along the length) gets the interpolated stress value.
            # This assumes stress varies only along the length, not circumferentially.
            # The segment_cnt > 0 check seems misplaced; coloring should probably always happen. Maybe avoids coloring the first point?
            # if segment_cnt > 0: # Let's remove this check, should always color
            for cnt,element in enumerate(ci): # Iterate through rows (length steps)
                   element[:] = C_dat[cnt] # Assign the interpolated stress to all points around circumference at this length step.
            # Transpose the color array to match the X, Y, Z data format for plot_surface.
            ci = np.transpose(ci)

            # --- Plot the Cylindrical Surface ---
            # Plot the surface using the calculated X, Y, Z points.
            # facecolors are determined by applying the colormap (cm.rainbow) to the normalized stress data (norm(ci)).
            # alpha sets the transparency.
            surf = ax1.plot_surface(x, y, z,facecolors=cm.rainbow(norm(ci)), alpha = 0.7)
            surf_list.append(surf) # Add the surface object to the list.

        # --- Final Plot Adjustments ---
        # Create a scalar mappable object for the color bar, using the same colormap.
        m = cm.ScalarMappable(cmap=cm.rainbow, norm=norm) # Use the same norm as the surface
        # Set the data range for the mappable (needed for the color bar).
        m.set_array(C_range) # Use the min/max stress values

        # Get the current axes instance (ax1).
        ax = plt.gca()

        # Add a color bar to the plot, linked to the scalar mappable 'm'. Shrink it slightly.
        col = plt.colorbar(m, shrink=.6, ax=ax)

        # Customize the color bar appearance.
        font_size = 14 # Font size for tick labels.
        col.ax.tick_params(labelsize=font_size) # Apply font size.
        col.set_label(label="Stress (PSI)", size = 18) # Set color bar label and size.

        # --- Set Plot Limits ---
        # Calculate the min/max coordinate values from the rotated data to set axis limits.
        max_x = 1*np.amax((coordinates[:,1].astype(float)))
        min_x = 1*np.amin((coordinates[:,1].astype(float)))
        max_y = 1*np.amax((coordinates[:,2].astype(float)))
        min_y = 1*np.amin((coordinates[:,2].astype(float)))
        max_z = 1*np.amax((coordinates[:,3].astype(float)))
        min_z = 1*np.amin((coordinates[:,3].astype(float)))

        # Find the absolute min/max across all coordinate dimensions for potentially equal aspect ratio.
        max_absolute = np.amax((coordinates[:,1:].astype(float)))
        min_absolute = np.amin((coordinates[:,1:].astype(float)))

        # Set the initial viewing angle for the 3D plot (elevation, azimuth).
        ax1.view_init(30,-45)

        # Set the limits for the X, Y, and Z axes based on the absolute min/max values.
        # This attempts to make the plot somewhat cubic, but doesn't guarantee equal aspect ratio.
        ax1.set_xlim(min_absolute,max_absolute)
        ax1.set_ylim(min_absolute,max_absolute)
        ax1.set_zlim(min_absolute,max_absolute) # Setting limits after view_init is usually fine.

        # --- Add Title Block Information using Tables in Subplots ---
        # Define grid spec coordinates for placing different parts of the title block at the bottom.
        # Each list is [y_min, y_max, x_min, x_max] in grid units (0-99).
        gs_list = [[-20,-13,-20,100], # Row 1 (Top of title block)
                   [-13,-10,-20,100], # Row 2
                   [-10,-7, -20,100], # Row 3
                   [-7, -3, -20,100], # Row 4 (Signature block)
                   [-3, 0,  -20,100]] # Row 5 (Bottom of title block) - Adjusted y_max to 0

        # Initialize a list to hold the text data for each part of the title block.
        array_list = []

        # Adjust subplot spacing to minimize gaps between title block elements.
        plt.subplots_adjust(wspace=0, hspace=0)

        # --- Prepare Title Block Text Data ---
        # Top rows of text (plant, system, plot type, load case). Reshape for table input.
        top_row = np.array(["GENERATING STATION",\
                            "HYDRAULIC SYSTEM MONITORING STATION",\
                            "ISOMETRIC STRESS CONTOUR - MECHANICAL",\
                            f"{loadcase} LOADCASE"]) # Use dynamic loadcase name
        top_row = top_row[:,np.newaxis] # Convert to column vector.
        array_list.append(top_row) # Add to list.

        # Second row text. Reshape for table input.
        second_row = np.array(["COMPANY NUCLEAR LLC"])[:,np.newaxis]
        array_list.append(second_row)

        # Third row text. Reshape for table input.
        third_row = np.array(["MECHANICAL DESIGN ENGINEERING","HANCOCKS BRIDGE, N.J."])[:,np.newaxis]
        array_list.append(third_row)

        # Fourth row: Signature block data (Prepared, Reviewed, Approved). Hardcoded names/dates.
        # This is a 2D array, directly usable by the table function.
        forth_row = np.array([["PREPARED BY:", "E. CASTILLO","SIGNATURE:","","DATE:",""],\
                              ["REVIEWED BY:", "E. APPIAH",  "SIGNATURE:","","DATE:",""],\
                              ["APPROVED BY:", "J. BOYER",    "SIGNATURE:","","DATE:",""]]) # Removed hardcoded numbers/dates
        array_list.append(forth_row)

        # Fifth row: Drawing/Figure title. Reshape for table input.
        fifth_row = np.array([f"FIGURE X - {loadcase} STRESS CONTOURS"])[:,np.newaxis] # Dynamic title
        array_list.append(fifth_row)

        # Define edge styles for each table in the title block ('open' = no border, 'closed' = border).
        edge_list = ["open","open","open","closed","open"] # Border around signature block.

        # --- Create Title Block Subplots and Tables ---
        # Iterate through the prepared text arrays and their corresponding grid spec positions.
        for index,array in enumerate(array_list):
            # Get the grid spec coordinates for the current subplot.
            y_min = gs_list[index][0]
            y_max = gs_list[index][1]
            x_min = gs_list[index][2]
            x_max = gs_list[index][3]

            # Add a new subplot (axes) at the specified grid location.
            k11 = fig1.add_subplot(gs[y_min:y_max,x_min:x_max])
            # Turn off ticks and labels for this subplot.
            k11.axes.get_xaxis().set_ticks([])
            k11.axes.get_yaxis().set_ticks([])
            # Remove padding around the subplot content.
            k11.axis('tight')
            # Turn off the axis frame.
            k11.axis('off') # Turn off axis lines as well

            # Create a table within the subplot using the text data.
            # loc='center', cellLoc='center' centers the table and text within cells.
            # edges sets the border style based on edge_list.
            k = k11.table(cellText=array, loc='center',cellLoc='center',edges=edge_list[index])

            # Adjust cell padding.
            k.auto_set_font_size(False) # Disable automatic font sizing.
            k.set_fontsize(8) # Set a fixed font size (adjust as needed).
            k.scale(1, 1.5) # Adjust cell height scaling if needed.

            # Special font size for certain rows (e.g., COMPANY NUCLEAR LLC, Figure title).
            if index in [1,4]: # Indices for second and fifth rows.
                 k.set_fontsize(10) # Slightly larger font size.
            if index == 3: # Signature block
                 k.set_fontsize(7) # Smaller font size


        # --- Save and Show Plot ---
        # Define the output filename for the plot PDF.
        plot_filename = f'runtime_files/stress_plot_{loadcase}.pdf' # Dynamic filename
        # Save the figure as a PDF file with high resolution (600 DPI).
        fig1.savefig(plot_filename, dpi=600)
        # Save as PNG (commented out).
        # fig1.savefig('rpic1.png', di=600)

        # Display the plot in the notebook output.
        plt.show()
        # Close the plot figure to free memory.
        plt.close(fig1)
        print(f"Stress plot saved to {plot_filename}")

    # Method to generate a 3D plot specifically for "THOR" data (elevations, etc.).
    # This method is largely a duplicate of get_stress_plot, with additions for a THOR data table.
    # Consider refactoring common plotting code into a helper function.
    def get_THOR_plot(self,loadcase):

        # Flag to track if the user provided an origin definition in the THOR comments.
        THOR_origin_given_by_user = False
        # Check if there was an error during THOR parameter parsing.
        if self.THOR_input_error:
            # If error occurred, print message and exit the method.
            print("THOR input is not in the correct format. Therefore, no THOR output can be generated.")
            return "" # Return empty string or None
        # If no THOR parsing error...
        else:
            # Get the rotated coordinates for the load case.
            coordinates = self.axes_rotation(loadcase,"hc") # Hardcoded plant "hc"

            # Convert the coordinate array into a dictionary keyed by node ID for easier lookup.
            # Stores coordinates as float arrays.
            self.THOR_coordinate_dic = {i[0]:i[1:].astype(float) for i in coordinates}

            # --- Find and Apply THOR Origin Offset ---
            # Iterate through the nodes found in the parsed THOR data dictionary.
            for node in self.THOR_node_dic:
                # Check if the 'ORIGIN' keyword exists within the parameters for this node.
                if "ORIGIN" in self.THOR_node_dic[node]:
                    # If found, assume this node defines the origin offset.
                    # print("ORIGIN NODE FOUND", node) # Debug print
                    # Extract the origin coordinate values from the THOR parameters.
                    # Assumes specific keys 'EAST_WEST_ORIGIN', 'UP_DOWN_ORIGIN', 'NORTH_SOUTH_ORIGIN'.
                    # These keys MUST match exactly what's parsed from the input file comments.
                    try: # Add error handling for missing keys
                        EAST_WEST_ORIGIN =   self.THOR_node_dic[node]['EAST_WEST_ORIGIN']
                        UP_DOWN_ORIGIN =     self.THOR_node_dic[node]['UP_DOWN_ORIGIN']
                        NORTH_SOUTH_ORIGIN = self.THOR_node_dic[node]['NORTH_SOUTH_ORIGIN']
                        # Set flag indicating origin was provided.
                        THOR_origin_given_by_user = True

                        # --- HARDCODED Coordinate Order Assumption ---
                        # Create the origin offset vector.
                        # CRITICAL ASSUMPTION: Assumes the order [EAST_WEST, UP_DOWN, NORTH_SOUTH] corresponds to the
                        # X, Y, Z axes *after* the axes_rotation method has been applied (which maps FEA Y to Plant Z, FEA Z to Plant -Y).
                        # This needs careful verification based on the specific plant coordinate system and FEA output.
                        # For 'hc' rotation (X=X_fea, Y=-Z_fea, Z=Y_fea):
                        # Plant X (East/West) likely corresponds to rotated X (FEA X).
                        # Plant Y (North/South) likely corresponds to rotated Y (FEA -Z).
                        # Plant Z (Up/Down) likely corresponds to rotated Z (FEA Y).
                        # Therefore, the vector should probably align with the rotated coordinates:
                        origin_vector = np.array([EAST_WEST_ORIGIN, NORTH_SOUTH_ORIGIN, UP_DOWN_ORIGIN]).astype(float)
                        # Let's assume the original hardcoded order was correct for the target system for now:
                        # origin_vector = np.array([EAST_WEST_ORIGIN,UP_DOWN_ORIGIN,NORTH_SOUTH_ORIGIN]).astype(float) # Original hardcoded order
                        print(f"Applying THOR origin offset from node {node}: {origin_vector}") # Info print
                        # Break the loop once the origin node is found and processed.
                        break # Assuming only one origin node definition
                    except KeyError as e:
                        print(f"Error: Missing expected ORIGIN coordinate key in THOR data for node {node}: {e}")
                        print("Cannot apply THOR origin offset.")
                        self.THOR_input_error = True # Set error flag as origin is incomplete
                        return "" # Exit method

            # --- Apply Origin Offset ---
            # If a valid origin was found and successfully parsed...
            if THOR_origin_given_by_user and not self.THOR_input_error:
                # Add the origin_vector offset to all coordinates in the THOR_coordinate_dic.
                # This translates the entire model relative to the specified plant origin.
                self.THOR_coordinate_dic = {node_id: self.THOR_coordinate_dic[node_id] + origin_vector
                                            for node_id in self.THOR_coordinate_dic}
                # print(self.THOR_coordinate_dic,"self.THOR_coordinate_dic") # Debug print
            elif not self.THOR_input_error: # If no origin was given, but no errors occurred
                 print("Warning: No THOR origin node definition found in input. Using relative coordinates.")
            # else: # Error occurred, message already printed

            # --- Plotting Setup ( Largely duplicated from get_stress_plot ) ---
            # Create figure and axes.
            fig1, _ = plt.subplots(1, 1, figsize=(22,17))
            # Define grid spec.
            gs = fig1.add_gridspec(100, 100)
            # Initialize list for surfaces (unused?).
            surf_list = []
            # Add main 3D subplot.
            ax1 = fig1.add_subplot(gs[1:-20,1:-20], projection='3d')
            # Turn off axes display.
            ax1._axis3don = False
            # Generate random vector (only used if plotting surfaces, which seems absent here).
            random_vec  = np.random.rand(3)
            random_vec /= np.linalg.norm(random_vec)

            print("Processing THOR Plot...") # Updated status message

            # --- Plot Geometry and Node Labels ---
            # Iterate through segments using tqdm progress bar.
            for segment_cnt in tqdm(range(len(self.node_connectivity_array))):
                # Get segment node IDs.
                segment = self.node_connectivity_array[segment_cnt]

                # Try to get coordinates from the potentially offset THOR_coordinate_dic.
                try:
                    # Get origin and destination coordinates using the potentially offset dictionary. Apply 1.3 scaling.
                    origin      = self.THOR_coordinate_dic[segment[0]] * 1.3
                    destination = self.THOR_coordinate_dic[segment[1]] * 1.3
                except KeyError as e:
                    # Handle cases where a node from connectivity might be missing in the coordinate dictionary.
                    print(f"Warning: Node {e} from connectivity not found in coordinate dictionary for THOR plot. Skipping segment {segment}.")
                    continue # Skip this segment

                # --- Plot Node Labels --- (Identical to get_stress_plot)
                node_font_size = 3
                if segment_cnt == 0:
                    bbox_props = dict(boxstyle="circle,pad=0.2", fc="w", ec="k", lw=.5)
                    t = ax1.text(origin[0], origin[1], origin[2], segment[0], ha="center", va="center", rotation=0,\
                                 size= node_font_size,\
                                 bbox=bbox_props,zorder=100)

                # Plot segment centerline (thin blue line).
                ax1.plot([origin[0],destination[0]], [origin[1],destination[1]], [origin[2],destination[2]], color = 'k',lw=.5, alpha=0.6) # Changed color/width

                if len(str(segment[1])) > 3:
                    label = str(segment[1]).replace(" ","")
                else:
                    label = str(segment[1])

                bbox_props = dict(boxstyle="circle,pad=0.2", fc="w", alpha=0.8, ec="k", lw=.5)
                t = ax1.text(destination[0], destination[1]+.03*(-1)**segment_cnt, destination[2], label,\
                             ha="center", va="center", rotation=0,\
                             size= node_font_size, bbox=bbox_props,zorder=100)

                # NOTE: The cylinder surface generation and plotting code from get_stress_plot is MISSING here.
                # This plot currently only shows centerlines and node labels. If surfaces are desired, copy that block here.

            # --- Plot Adjustments (Limits, View) --- (Identical to get_stress_plot, but uses THOR_coordinate_dic)
            # Get all coordinate values (X, Y, Z) from the potentially offset dictionary.
            all_coords_flat = np.array(list(self.THOR_coordinate_dic.values())).flatten()
            if len(all_coords_flat) > 0: # Check if there are any coordinates
                max_absolute = np.amax(all_coords_flat)
                min_absolute = np.amin(all_coords_flat)
            else: # Default limits if no coordinates
                max_absolute = 10
                min_absolute = -10

            ax1.view_init(30,-45) # Set view angle.

            # Set axis limits using the range of potentially offset coordinates.
            ax1.set_xlim(min_absolute,max_absolute)
            ax1.set_ylim(min_absolute,max_absolute)
            ax1.set_zlim(min_absolute,max_absolute)

            # --- Add Title Block (Similar to get_stress_plot, but modified text) ---
            # Grid spec positions for title block elements.
            gs_list = [[-20,-13,-20,100],[-13,-10,-20,100],[-10,-7,-20,100],[-7,-3,-20,100],[-3,0,-20,100]] # Adjusted y_max for last row
            # List to hold text arrays for title block.
            array_list = []
            # Adjust subplot spacing.
            plt.subplots_adjust(wspace=0, hspace=0)

            # --- Prepare Title Block Text --- (Modified for THOR)
            top_row = np.array(["GENERATING STATION",\
                                "HYDRAULIC SYSTEM MONITORING STATION",\
                                "ISOMETRIC VIEW - THOR NODES",\ # Modified title
                                f"{loadcase} LOADCASE"])
            top_row = top_row[:,np.newaxis]
            array_list.append(top_row)

            second_row = np.array(["COMPANY NUCLEAR LLC"])[:,np.newaxis]
            array_list.append(second_row)

            third_row = np.array(["MECHANICAL DESIGN ENGINEERING","SITE LOCATION"])[:,np.newaxis]
            array_list.append(third_row)

            # Signature block (same as before).
            forth_row = np.array([["PREPARED BY:", "E. CASTILLO","SIGNATURE:","","DATE:",""],\
                                  ["REVIEWED BY:", "REVIEWER NAME",  "SIGNATURE:","","DATE:",""],\
                                  ["APPROVED BY:", "APPROVER NAME",    "SIGNATURE:","","DATE:",""]])
            array_list.append(forth_row)

            # Bottom title row (modified for THOR).
            fifth_row = np.array(["FIGURE Y - THOR NODE LOCATIONS & ELEVATIONS"])[:,np.newaxis] # Modified title
            array_list.append(fifth_row)

            # Edge styles for title block tables.
            edge_list = ["open","open","open","closed","open"]

            # --- Prepare THOR Data Table ---
            # Define headers for the THOR data table.
            THOR_headers = np.array(["ME101 NODE NUMBER","THOR NODE NUMBER","ELEVATION (FT)"])

            # Create the data array for the THOR table.
            # Iterates through coordinate dictionary, checks if node exists in THOR data, extracts relevant info.
            # HARDCODED: Assumes elevation is the 3rd element (index 2) of the coordinate vector in self.THOR_coordinate_dic.
            # This depends critically on the output of axes_rotation and the application of the origin offset.
            # If axes_rotation puts Plant Z (Up/Down) in index 2, this is correct.
            THOR_data_list = []
            for node_id in self.THOR_coordinate_dic:
                if node_id in self.THOR_node_dic: # Check if this node has associated THOR parameters
                   try: # Add error handling for missing THOR NODE NUMBER
                       thor_node_number = self.THOR_node_dic[node_id]["THOR NODE NUMBER"]
                       elevation = self.THOR_coordinate_dic[node_id][2] # Assumes Z (elevation) is index 2
                       THOR_data_list.append([node_id, thor_node_number, np.around(elevation, decimals=3)])
                   except KeyError:
                       print(f"Warning: Missing 'THOR NODE NUMBER' parameter for ME101 node {node_id}.")
                   except IndexError:
                        print(f"Warning: Coordinate data for node {node_id} has fewer than 3 dimensions.")

            # Convert the list of data rows to a NumPy array.
            THOR_data = np.asarray(THOR_data_list)

            # Combine headers and data into the final table array. Check if data exists.
            if THOR_data.shape[0] > 0: # Check if any THOR data was successfully extracted
                THOR_table = np.vstack((THOR_headers,THOR_data))
            else:
                THOR_table = np.array([THOR_headers, ["No THOR data found", "", ""]]) # Placeholder if no data
                print("Warning: No valid THOR node data found to display in the table.")


            # --- Create Subplot for THOR Table ---
            # Define grid spec position for the THOR table (top right area, distinct from title block).
            y_min_table = 5 # Top edge
            y_max_table = 75 # Bottom edge
            x_min_table = 65 # Left edge (Adjust position as needed)
            x_max_table = 98 # Right edge

            # Add subplot for the table title.
            # k13 = fig1.add_subplot(gs[y_min_table-5:y_min_table, x_min_table:x_max_table]) # Position title above table
            # k13.set_title('THOR NODE ELEVATION TABLE',fontsize= 12) # Set title
            # k13.axis('off') # Turn off axes for title subplot

            # Add subplot where the table will be placed.
            k12 = fig1.add_subplot(gs[y_min_table:y_max_table, x_min_table:x_max_table])
            # Configure axes for the table subplot.
            k12.axes.get_xaxis().set_ticks([]) # No ticks
            k12.axes.get_yaxis().set_ticks([])
            k12.axis('tight') # Remove padding
            k12.axis('off') # Turn off frame/axes lines

            # Add the THOR data table to this subplot.
            # loc='center' places the table in the center of the subplot bounds. edges='closed' adds borders.
            k_thor = k12.table(cellText=THOR_table, loc='center', cellLoc='center', edges="closed", colWidths=[0.3, 0.3, 0.3]) # Adjust colWidths as needed
            k_thor.auto_set_font_size(False) # Manual font size
            k_thor.set_fontsize(8) # Set font size
            k_thor.scale(1, 1.5) # Scale cell height if needed

            # --- Create Title Block Subplots and Tables (Identical code as before) ---
            # Iterate through the title block text arrays.
            for index,array in enumerate(array_list):
                # Get grid spec coordinates.
                y_min = gs_list[index][0]
                y_max = gs_list[index][1]
                x_min = gs_list[index][2]
                x_max = gs_list[index][3]
                # Add subplot.
                k11 = fig1.add_subplot(gs[y_min:y_max,x_min:x_max])
                # Configure axes.
                k11.axes.get_xaxis().set_ticks([])
                k11.axes.get_yaxis().set_ticks([])
                k11.axis('tight')
                k11.axis('off') # Turn off frame
                # Add table.
                k = k11.table(cellText=array, loc='center',cellLoc='center',edges=edge_list[index])
                # Configure table font/padding.
                k.auto_set_font_size(False)
                k.set_fontsize(8)
                k.scale(1, 1.5)
                if index in [1,4]: # COMPANY / Figure Title rows
                    k.set_fontsize(10)
                if index == 3: # Signature block
                    k.set_fontsize(7)

            # --- Save Plot and THOR Data ---
            # Define plot output filename.
            thor_plot_filename = 'runtime_files/THOR_plot.pdf' # Changed path to runtime_files
            # Save the figure as PDF.
            fig1.savefig(thor_plot_filename, dpi=600)
            # Show the plot.
            plt.show()
            # Close the figure.
            plt.close(fig1)
            print(f"THOR plot saved to {thor_plot_filename}")

            # --- Write THOR Coordinates to CSV ---
            # Define CSV output filename.
            thor_csv_filename = 'runtime_files/THOR_COORDINATES.csv' # Changed path

            # Open the CSV file in write mode ('w') to overwrite existing content.
            with open(thor_csv_filename,'w') as file:

                # Write the header row.
                file.write("ME101 NODE NUMBER,THOR NODE NUMBER,NODE ELEVATION (FT)")
                
            # Open the CSV file in append mode ('a') to add data rows.
            with open(thor_csv_filename,'a') as file:

                 # Iterate through the successfully processed THOR data list.
                for row_data in THOR_data_list:

                    # Format the data as a comma-separated string. Ensure elevation is string.
                    current_string = ",".join(map(str, row_data)) # Convert all elements to string

                    # Write the data row, preceded by a newline character.
                    file.write("\n" + current_string)

            print(f"THOR coordinates saved to {thor_csv_filename}")


    # Method to parse the text content of the Pipe Stress Summary section.
    def parse_pipe_stress_summary(self):

        # Get the start and end line numbers previously identified for the summary section. Check if found.
        if self.pipe_stress_summary_start_line_number is None or self.pipe_stress_summary_end_line_number is None:
            print("Warning: Pipe stress summary section boundaries not found. Cannot parse summary text.")
            self.stress_summary_list = ["Error: Pipe stress summary not found in output file."]
            self.stress_summary_text = "Error: Pipe stress summary not found."
            self.stress_ratio_list = []
            return

        # Define how many lines *before* the header to include (often contains titles).
        leading_lines = 16 # Assumed number of lines for context/title before header.

        # Define how many lines *after* the end keyword to include (often contains footers/signatures).
        extra_end_lines = 2 # Assumed number of lines after the end keyword.

        # Calculate the actual start line index, ensuring it's not negative.
        start_index = max(0, self.pipe_stress_summary_start_line_number - leading_lines)
        # Calculate the actual end line index.
        end_index = self.pipe_stress_summary_end_line_number + extra_end_lines

        # Extract the relevant lines from the raw data list.
        stress_summary_list_raw = self.raw_data_list[start_index : end_index] # Slice includes start, excludes end+1

        # Clean the extracted lines: remove newline characters and potentially problematic control characters (\x01).
        self.stress_summary_list = [i.replace("\n","").replace("\x01"," ") for i in stress_summary_list_raw]

        # --- Extract Stress Ratios ---

        # Initialize a list to store extracted stress ratios (Actual Stress / Allowable Stress).
        self.stress_ratio_list = []

        # Iterate through the cleaned lines of the summary section.
        for line in self.stress_summary_list:
            # Extract the content from columns 72-79, assumed to contain the stress ratio. Strip whitespace and remove '*'.
            stress_ratio_str = line[72:79].strip().replace("*","") if len(line) >= 79 else "" # Add length check

            # Check if the extracted string represents a valid number (allowing decimals).
            # Use try-except for robust float conversion.
            try:
                stress_ratio_val = float(stress_ratio_str)
                # If conversion is successful, append the float value to the list.
                self.stress_ratio_list.append(stress_ratio_val)

            except ValueError:
                # If conversion fails, it's not a valid number in this column, so ignore it.
                pass

        # --- Determine Summary Text based on Ratios ---
        # Check if any stress ratios were found.
        if not self.stress_ratio_list: # If the list is empty
             self.stress_summary_text = "The pipe stress summary is shown below. No valid stress ratios were found in the expected columns (72-79)."
        
        # Check if the maximum stress ratio found is less than 1.0.
        elif max(self.stress_ratio_list) < 1.0:
            
            # If all ratios are < 1.0, the stresses meet the code requirements.
            self.stress_summary_text = "The pipe stress summary is shown below, with all actual pipe stresses meeting code allowable stresses."
        else:
            
            # If any ratio is >= 1.0, allowable stresses are exceeded.
            self.stress_summary_text = f"The pipe stress summary is shown below, with code allowable stresses potentially being exceeded. " \
                                      f"The maximum stress ratio found is {max(self.stress_ratio_list):.3f}." # Format max ratio

    # Method to populate a DOCX template file with parsed data.
    def parse_text(self):

        # Construct the path to the combined template file created by compose_document.
        template_path = "output_files/combined.docx" # Assumes compose_document was run first
        
        try:
            # Open the combined DOCX template file.
            document = docx.Document(template_path)
        
        except Exception as e:
            print(f"Error opening template file '{template_path}': {e}")
            print("Ensure 'compose_document()' was executed successfully first.")
            return # Exit if template cannot be opened

        # --- Set Default Font Style ---
        
        # Get the default 'Normal' style object.
        style = document.styles['Normal']
       
        # Access the font object within the style.
        font = style.font
        
        # Set the default font name.
        font.name  = 'Arial'
        
        # Set the default font size.
        font.size  = Pt(9)

        # --- Define Style for Added/Modified Text ---
        # Get the 'Heading 2' style (or choose another appropriate style).
        # This style will be used to format the text inserted by the script.
        # Using a heading style might have unintended formatting consequences; 'Normal' might be safer,
        # or define a custom style in the template. Let's assume 'Heading 2' is intentional for now.
        style_added_text = document.styles['Normal'] # Changed to Normal, modify if needed
        font_added_text = style_added_text.font
        font_added_text.name  = 'Arial' # Consistent font
        font_added_text.size  = Pt(9)   # Consistent size
        
        # Set the color for added text (a specific shade of blue).
        font_added_text.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        # font_added_text.bold = True # Optionally make added text bold


        # Get all sections in the document (usually one, but can be more).
        sections = document.sections
        
        # Get the count of sections.
        section_count = len(sections)

        # --- Process Headers ---
        # Loop through each section to process its header.
        for section_number in range(section_count):
            # Get the current section object.
            section = sections[section_number]
            # Get the header associated with this section.
            header = section.header
            # Find all tables within the header.
            header_tables = header.tables
            # Loop through each table in the header.
            for table in header_tables :
                # Loop through each row in the table.
                for row in table.rows:
                    # Loop through each cell in the row.
                    for cell in row.cells:
                        # Loop through each paragraph within the cell.
                        # Placeholders are expected to be within paragraphs.
                        for paragraph in cell.paragraphs:
                            # --- Placeholder Replacement Logic (Header) ---
                            # Find the start and end markers for placeholders (e.g., <PLACEHOLDER>).
                            start = paragraph.text.find("<")
                            end = paragraph.text.find(">")

                            # Initialize variables for concatenated text and identified parameters.
                            concat_parameters = str("")
                            parameter_list_in_para = [] # Use different name to avoid conflict

                            # If both start and end markers are found in the paragraph text...
                            if start != -1 and end != -1:
                                # Identify all known coversheet parameter keys present within the placeholder markers.
                                # This allows multiple parameters like "<PARAM1>-<PARAM2>" in one cell.
                                # It assumes parameter names don't contain '<' or '>'.
                                full_placeholder_text = paragraph.text[start:end+1] # Get text between and including <>
                                parameter_list_in_para = [key for key in self.coversheet_parameter_list if key in full_placeholder_text]

                                # --- Assemble Text for Header Cell ---
                                # If exactly one parameter key was found...
                                if len(parameter_list_in_para) == 1:
                                    parameter = parameter_list_in_para[0]
                                    # Get the value from the coversheet dictionary, handle missing keys gracefully.
                                    concat_parameters = self.coversheet_parameter_dic.get(parameter, f"<{parameter}_NOT_FOUND>")
                                # If more than one parameter key was found...
                                elif len(parameter_list_in_para) > 1:
                                    # Concatenate the values, potentially adding prefixes/suffixes.
                                    # This logic assumes parameters are meant to be joined, e.g., "<DOC_NUM>-<REV>".
                                    # The order might depend on the order in self.coversheet_parameter_list, which might not match the placeholder.
                                    # A more robust approach would parse the exact structure within <>.
                                    temp_text = ""
                                    for parameter in parameter_list_in_para:
                                        add_text = str(self.coversheet_parameter_dic.get(parameter, f"<{parameter}_NOT_FOUND>"))
                                        # Special handling for revision numbers (add "R").
                                        if "REVISION" in parameter.upper(): # Case-insensitive check
                                            add_text = "R"+add_text
                                        temp_text += add_text + "-" # Add separator (adjust if needed)
                                    concat_parameters = temp_text.rstrip("-") # Remove trailing separator
                                else: # No known parameter found within <>
                                     concat_parameters = full_placeholder_text # Keep the original placeholder text

                                # --- Replace Placeholder with Value ---
                                # Clear the original paragraph text.
                                paragraph.text = ""
                                # Add a new run containing the assembled text.
                                run = paragraph.add_run(concat_parameters)
                                # Apply the defined style (color, font) to the added run.
                                run.font.name = font_added_text.name
                                run.font.size = font_added_text.size
                                run.font.color.rgb = font_added_text.color.rgb
                                # run.bold = font_added_text.bold # Apply bold if set in style

        # --- Process Main Document Body (Tables and Paragraphs) ---
        # Get all tables in the main body of the document.
        tables = document.tables
        # Counter for iterating through multi-value 'COMPONENT' parameters.
        component_count = 0
        # Loop through each table in the document body.
        for table in tables:
            # Loop through each row in the table.
            for row in table.rows:
                # Loop through each cell in the row.
                for cell in row.cells:
                    # Loop through each paragraph in the cell.
                    for paragraph in cell.paragraphs:
                        # Check if the paragraph text contains placeholder markers.
                        if "<" in paragraph.text and ">" in paragraph.text:
                            # Extract the text within the markers, assuming it's a single key.
                            # This simple extraction might fail for complex placeholders like "<PARAM1>-<PARAM2>".
                            parameter = paragraph.text.strip().replace("<","").replace(">","")

                            # --- Placeholder Replacement Logic (Body Tables) ---
                            # Check if the extracted parameter is a known coversheet parameter.
                            if parameter in self.coversheet_parameter_list:
                                # Clear the paragraph.
                                paragraph.text = ""
                                # Add run with the value from the dictionary (handle missing keys).
                                run = paragraph.add_run(self.coversheet_parameter_dic.get(parameter, f"<{parameter}_NOT_FOUND>"))
                                # Apply style.
                                run.font.name = font_added_text.name
                                run.font.size = font_added_text.size
                                run.font.color.rgb = font_added_text.color.rgb
                                # run.bold = font_added_text.bold

                            # Handle the special 'COMPONENT' parameter (multi-value).
                            elif parameter == "COMPONENT":
                                # Get the list of components from the dictionary.
                                component_list = self.coversheet_parameter_dic.get(parameter, [])
                                # Check if there are components left to insert.
                                # The 'or True' bypasses the limit check - remove 'or True' to enforce limit.
                                if component_count < len(component_list):# and (len(component_list) <= 10 or True):
                                    # Clear the paragraph.
                                    paragraph.text = ""
                                    # Add run with the next component from the list.
                                    run = paragraph.add_run(component_list[component_count])
                                    # Apply style.
                                    run.font.name = font_added_text.name
                                    run.font.size = font_added_text.size
                                    run.font.color.rgb = font_added_text.color.rgb
                                    # run.bold = font_added_text.bold
                                    # Increment the component counter.
                                    component_count +=1
                                else: # If no more components or limit reached
                                     paragraph.text = "" # Clear the placeholder if no component to add

                            # Handle parameters defined as long text (potentially multi-line).
                            elif parameter in self.long_text_parameter_list:
                                # Clear the paragraph.
                                paragraph.text = ""
                                # Add run with the (concatenated) value from the dictionary.
                                run = paragraph.add_run(self.coversheet_parameter_dic.get(parameter, f"<{parameter}_NOT_FOUND>"))
                                # Apply style.
                                run.font.name = font_added_text.name
                                run.font.size = font_added_text.size
                                run.font.color.rgb = font_added_text.color.rgb
                                # run.bold = font_added_text.bold

                            # Handle parameters related to nozzle load check results.
                            # Assumes self.nozzle_load_parameter_dic is populated by nozzle_load_check().
                            elif parameter in self.nozzle_analysis_parameter_list:
                                # Check if the dictionary exists.
                                if hasattr(self, 'nozzle_load_parameter_dic'):
                                    # Clear the paragraph.
                                    paragraph.text = ""
                                    # Get the value, convert to string, handle missing keys.
                                    value_str = str(self.nozzle_load_parameter_dic.get(parameter, f"<{parameter}_NOT_FOUND>"))
                                    run = paragraph.add_run(value_str)
                                    # Apply style.
                                    run.font.name = font_added_text.name
                                    run.font.size = font_added_text.size
                                    run.font.color.rgb = font_added_text.color.rgb
                                    # run.bold = font_added_text.bold
                                else:
                                    paragraph.text = f"<{parameter}_DICT_MISSING>" # Error message
                            # else: # Parameter not recognized, leave placeholder as is
                               # paragraph.text = f"<{parameter}_UNKNOWN>" # Optional: Mark unknown placeholders

        # --- Process Paragraphs in Main Document Body ---
        # Get all paragraphs in the main body (outside tables).
        all_paragraphs = document.paragraphs
        # Iterate through each paragraph.
        for each_paragraph in all_paragraphs:
            # Check for placeholder markers.
            if "<" in each_paragraph.text and ">" in each_paragraph.text:
                # Extract the parameter name (assuming simple placeholder).
                parameter = each_paragraph.text.strip().replace("<","").replace(">","")

                # --- Handle Special Paragraph Placeholders ---
                # Check if it's a parameter designated for Courier font formatting (e.g., Pipe Stress Summary).
                if parameter in self.courier_long_text_parameter_dic:
                    # Ensure the stress summary data has been parsed.
                    if not hasattr(self, 'stress_summary_list') or not hasattr(self, 'stress_summary_text'):
                         print(f"Warning: Stress summary data not available for placeholder <{parameter}>.")
                         each_paragraph.text = f"<{parameter}_DATA_MISSING>"
                         continue # Skip this placeholder

                    # --- Insert Stress Summary Section ---
                    # Construct the section heading using the parameter name and its assigned section number.
                    text_heading = self.courier_long_text_parameter_dic[parameter] + " " + parameter
                    # Insert the heading paragraph *before* the current placeholder paragraph.
                    # Apply a distinct style (e.g., 'Heading 1') if desired, otherwise uses default.
                    # Use the 'style_added_text' for consistency or define a specific heading style.
                    heading_para = each_paragraph.insert_paragraph_before(text_heading) # Apply style later if needed
                    # Optionally apply formatting directly to the heading run
                    # heading_para.runs[0].font.bold = True
                    # heading_para.runs[0].font.size = Pt(10)

                    # Insert the summary text paragraph (pass/fail statement) before the placeholder.
                    # Apply the defined 'added text' style.
                    summary_text_para = each_paragraph.insert_paragraph_before(self.stress_summary_text, style=style_added_text.name) # Use style name
                    # Insert an empty paragraph for spacing before the detailed summary.
                    each_paragraph.insert_paragraph_before("")

                    # Clear the original placeholder paragraph text.
                    each_paragraph.text = ""

                    # Insert the detailed pipe stress summary lines.
                    # Iterate through the stored summary lines.
                    for line in self.stress_summary_list:
                        # Add the line as a new run to the (now empty) original placeholder paragraph.
                        run = each_paragraph.add_run(line)
                        # Apply specific formatting for the summary (Courier New, small size, blue color).
                        run.font.name = "Courier New"
                        run.font.size = Pt(6)
                        run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
                        # Add a line break after each line of the summary.
                        run.add_break() # Use line break instead of new paragraph for better block formatting.

                    # Insert an empty paragraph after the summary block for spacing.
                    each_paragraph.insert_paragraph_before("") # This inserts *before*, so add after loop if needed.
                    # Let's add space *after* by adding a new paragraph:
                    document.add_paragraph("") # Adds paragraph at the end of the document, not ideal.
                    # Alternative: Insert before the *next* paragraph if possible, or rely on template spacing.

                # elif parameter in self. ... : # Handle other paragraph-level placeholders if any
                #     pass

        # --- Save the Modified Document ---
        # Define the final output filename.
        output_doc_path = 'output_files/OUTPUT_MERGED.docx'
        try:
            # Save the document with all replacements applied.
            document.save(output_doc_path)
            print(f"Successfully processed and saved document to {output_doc_path}")
        except Exception as e:
            print(f"Error saving final document '{output_doc_path}': {e}")
        # Optional: Copy to another location (e.g., Dropbox) - commented out.
        # copyfile("me101/OUTPUT_MERGED.docx","Dropbox/me101/OUTPUT_MERGED.docx")


    # Method to perform nozzle load calculations based on parsed geometry/parameters.
    def nozzle_load_check(self):

        # --- Configuration ---
        # List of node IDs designated as nozzle points to be checked. Hardcoded.
        self.nozzle_load_check_node_list = ["50"] # Example: Check node "50"
        # Material yield strength (PSI). Hardcoded - should ideally be read from input or config.
        self.yield_strength = 35000
        # Factors K1 and K2 used in allowable load calculations. Hardcoded - based on specific code/standard?
        self.k1 = 0.01 # Factor for axial/shear allowable.
        self.k2 = 0.1 # Factor for bending allowable.

        # Initialize dictionary to store calculated nozzle parameters and allowables for the *first* nozzle checked.
        # This method currently only processes the first node in nozzle_load_check_node_list effectively,
        # as it overwrites self.nozzle_load_parameter_dic in each loop iteration.
        # Needs modification to handle multiple nozzles if required.
        self.nozzle_load_parameter_dic = {}

        # Iterate through the list of designated nozzle node IDs.
        for nozzle_check_point in self.nozzle_load_check_node_list:

            # Find all entries in the raw segment parameter list that correspond to the current nozzle node ID.
            # This list contains pairs [node_id, {param_dict}]
            check_node_list = [node for node in self.raw_segment_parameter_list if node[0]==nozzle_check_point]

            # Check if the nozzle node was found in the parsed geometry data.
            if not check_node_list:
                print(f"Warning: Nozzle check point '{nozzle_check_point}' not found in parsed segment data. Skipping.")
                continue # Skip to the next nozzle node ID

            # --- Determine Nozzle Properties ---
            # Find the segment entry corresponding to the minimum wall thickness reported at this node.
            # Assumes 'THI' (thickness) is present and numeric in the parameter dictionary for each entry.
            try:
                # Calculate index of the minimum thickness entry.
                node_min_thickness_index = np.argmin([float(node[1]["THI"]) for node in check_node_list])
                # Get the parameter dictionary for the entry with minimum thickness.
                min_thi_params = check_node_list[node_min_thickness_index][1]
            except (KeyError, ValueError, IndexError) as e:
                print(f"Warning: Could not determine minimum thickness for nozzle '{nozzle_check_point}'. Missing or invalid 'THI' parameter? Error: {e}. Skipping.")
                continue # Skip to the next nozzle node ID

            # Store the nozzle node ID.
            self.nozzle_load_parameter_dic['NOZZLE NODE POINT'] =  nozzle_check_point

            # --- Extract and Calculate Geometric Properties ---
            try:
                # Get Outer Diameter (OD) from the min thickness entry's parameters. Convert to float.
                pipe_OD = float(min_thi_params["OD"])
                # Store OD, rounded to 3 decimal places.
                self.nozzle_load_parameter_dic['OUTSIDE DIAMETER'] = np.around(pipe_OD,decimals=3)

                # Get Thickness (THI) from the min thickness entry's parameters. Convert to float.
                pipe_THICKNESS = float(min_thi_params["THI"]) # Typo fixed: THICKNESS
                # Store Thickness, rounded to 3 decimal places.
                self.nozzle_load_parameter_dic['NOZZLE THICKNESS'] = np.around(pipe_THICKNESS,decimals=3)

                # Calculate Inner Diameter (ID).
                pipe_ID = pipe_OD - 2 * pipe_THICKNESS
                # Calculate cross-sectional area of the pipe wall.
                pipe_area = np.pi*(  (pipe_OD/2)**2 - (pipe_ID/2)**2  )
                # Store Area, rounded to 3 decimal places.
                self.nozzle_load_parameter_dic['NOZZLE AREA'] = np.around(pipe_area,decimals=3)

                # Calculate the Section Modulus (Z) of the pipe.
                # Formula: Z = pi * (Do^4 - Di^4) / (32 * Do)
                pipe_SECTION_MODULUS = np.pi*(pipe_OD**4 - pipe_ID**4)/(32*pipe_OD)
                # Store Section Modulus, rounded to 5 decimal places.
                self.nozzle_load_parameter_dic['NOZZLE SECTION MODULUS'] = np.around(pipe_SECTION_MODULUS,decimals=5)

            except (KeyError, ValueError) as e:
                 print(f"Warning: Missing or invalid 'OD' or 'THI' for nozzle '{nozzle_check_point}'. Error: {e}. Skipping calculations.")
                 # Clear potentially partially filled dict for this nozzle
                 self.nozzle_load_parameter_dic = {k: v for k, v in self.nozzle_load_parameter_dic.items() if k == 'NOZZLE NODE POINT'}
                 continue # Skip calculation part

            # --- Store Material/Factor Properties ---
            # Store Yield Strength (rounded to integer).
            self.nozzle_load_parameter_dic['YIELD'] = np.around(self.yield_strength,decimals = 0)
            # Store K1 factor.
            self.nozzle_load_parameter_dic['K1'] = self.k1
            # Store K2 factor.
            self.nozzle_load_parameter_dic['K2'] = self.k2

            # --- Calculate Allowable Loads ---
            # Calculate Allowable Axial Load = K1 * Sy * Area. Convert to integer.
            self.nozzle_load_parameter_dic['ALLOWABLE AXIAL LOAD'] = int(self.k1 * self.yield_strength * pipe_area)
            # Assume Allowable Resultant Shear Load is the same as Axial Load. This might be specific to the code/standard used.
            self.nozzle_load_parameter_dic['ALLOWABLE RESULTANT SHEAR LOAD'] = int(self.nozzle_load_parameter_dic['ALLOWABLE AXIAL LOAD'])
            # Calculate Allowable Bending Moment = K2 * Sy * Z. Convert to integer. Divide by 12 (for ft-lbs?). Units need verification.
            self.nozzle_load_parameter_dic['ALLOWABLE BENDING MOMENT'] = int(self.k2 * self.yield_strength * \
                                                                            pipe_SECTION_MODULUS / 12) # Division by 12 suggests Z is in^3, result in ft-lb?
            # Assume Allowable Torsion Moment is twice the Bending Moment. Specific to code/standard?
            self.nozzle_load_parameter_dic['ALLOWABLE TORSION MOMENT'] = int(2 * self.nozzle_load_parameter_dic['ALLOWABLE BENDING MOMENT'])

            # NOTE: Since this dictionary is overwritten in each loop iteration, only the results
            # for the *last* nozzle ID in nozzle_load_check_node_list will be available
            # in self.nozzle_load_parameter_dic after the loop finishes.
            # If results for all nozzles are needed, store them in a list or a nested dictionary.
            print(f"Nozzle check calculations completed for node: {nozzle_check_point}")


    # Method to combine multiple DOCX files into one.
    def compose_document(self):
        # Define the path to the primary template file (cover sheet).
        master_path = "runtime_files/cover.docx"
        # Define the path to the secondary template file (nozzle loads).
        doc2_path = "runtime_files/nozzle_loads.docx"
        # Define the output path for the combined document.
        output_path = "output_files/combined.docx"

        try:
            # Open the master document.
            master = docx.Document(master_path)
            # Add a page break after the content of the master document.
            master.add_page_break()
            # Create a Composer object, initialized with the master document.
            composer = Composer(master)
            # Open the second document to be appended.
            doc2 = docx.Document(doc2_path)
            # Append the content of the second document to the composer object.
            composer.append(doc2)
            # Save the combined document to the specified output path.
            composer.save(output_path)
            print(f"Successfully composed document: {output_path}")
        except Exception as e:
            print(f"Error composing document from {master_path} and {doc2_path}: {e}")



# Record the starting time for performance measurement.
start_time = time.time() # Use a different variable name

# --- Initialize Parameters ---
# List of keywords identifying relevant deadweight load cases in the output file.
data_echo_deadweight_keywords = ["WTDW","DW"]
# Path to the main FEA output file to be processed.
filename = "input_files/out.txt"
# Maximum number of consecutive non-loadcase lines to scan before stopping load case search.
max_blank_line_cnt = 500

# --- Instantiate the Class ---
# Create an instance of the piping_system class, passing the initialization parameters.
hcu = piping_system(filename,data_echo_deadweight_keywords,max_blank_line_cnt)

# --- Execute Processing Steps ---
# 1. Parse the raw FEA output file to identify load cases.
hcu.parse_load_cases() # This also calls parse_raw_data if needed.

# 2. Perform nozzle load calculations based on parsed geometry and hardcoded parameters.
#    This populates self.nozzle_load_parameter_dic (likely only for the last nozzle in the list).
hcu.nozzle_load_check()

# 3. Parse the text content and stress ratios from the Pipe Stress Summary section.
#    This populates self.stress_summary_list and self.stress_summary_text.
hcu.parse_pipe_stress_summary()

# 4. Combine the cover sheet and nozzle load templates into a single document.
#    Creates 'output_files/combined.docx'.
hcu.compose_document()

# 5. Parse the combined template ('combined.docx') and replace placeholders with data
#    parsed earlier (coversheet info, nozzle calcs, stress summary text).
#    Saves the final report as 'output_files/OUTPUT_MERGED.docx'.
hcu.parse_text()

# 6. Check coordinate consistency across load cases (optional, prints message).
# hcu.get_cordinate_distance() # Commented out

# 7. Parse node connectivity (segments) from the stress analysis section.
#    Populates self.node_connectivity_array.
hcu.get_node_connectivities()

# 8. Parse node coordinates for each load case from the NODE DATA sections.
#    Populates self.loadcase_geometry_dic and self.all_coordinate_id_list.
node_coordinates = hcu.parse_node_coordinates() # Return value assigned but also stored in self.loadcase_geometry_dic

# 9. Parse sustained stress values from the stress analysis section.
#    Populates self.sustained_stresses.
hcu.parse_stresses()

# 10. Generate and save/show the 3D stress contour plot for a specific load case ("WTDW").
#     Commented out call for "WT01".
# hcu.get_stress_plot("WT01")
hcu.get_stress_plot("WTDW") # Generate plot for WTDW load case.

# 11. Generate and save/show the 3D THOR plot (if THOR data was parsed correctly).
#     This plots geometry, labels, and a table of THOR node elevations.
# hcu.get_THOR_plot("WTDW") # Commented out

# Debug print (commented out).
# print(node_coordinates['WTDW'])

# --- Final Timing ---
# Calculate and print the total execution time.
print("Total Execution Time:",time.time()-start_time) # Use start_time variable



